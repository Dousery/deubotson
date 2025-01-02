from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.common.by import By
from docx import Document

headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/87.0.4280.88 Safari/537.36"
}
driver = webdriver.Chrome()

def get_ogretim_uyeleri(url):
    driver.get(url)
    soup = BeautifulSoup(driver.page_source, "html.parser")
    section = soup.find("section", class_="blog-listing")

    ogretim_uyeleri = []
    for index, hoca in enumerate(section.find_all("div", class_="su-row")):
        if index == 0:  # İlk öğeyi atlamak için
            continue
        
        isim = hoca.find("strong").get_text(strip=True)
        unvan = hoca.find("em").get_text(strip=True)

        telefon_element = hoca.find("strong", string="Tel:")
        if telefon_element and telefon_element.next_sibling:
            telefon = telefon_element.next_sibling.strip()
        else:
            telefon = None

        email_element = hoca.find("a", href=lambda href: href and "mailto:" in href)
        email = email_element.get_text(strip=True) if email_element else None

        # Araştırma alanlarını al
        alanlar_element = hoca.find("div", class_="su-column su-column-size-2-5")
        if alanlar_element:
            alanlar = alanlar_element.get_text(strip=True)
        else:
            alanlar = "Araştırma alanı bulunamadı"

        ogretim_uyeleri.append({
            "İsim": isim,
            "Unvan": unvan,
            "Telefon": telefon,
            "E-posta": email,
            "Araştırma Alanları": alanlar
        })

    word_path = 'scraping.docx'
    doc = Document(word_path)

    doc.add_heading('Öğretim Üyeleri', level=1)
    # Verileri ekle
    for uye in ogretim_uyeleri:
        doc.add_heading(uye['İsim'], level=2)
        doc.add_paragraph(f"Unvan: {uye['Unvan']}")
        doc.add_paragraph(f"Telefon: {uye['Telefon']}")
        doc.add_paragraph(f"E-posta: {uye['E-posta']}")
        doc.add_paragraph(f"Araştırma Alanları: {uye['Araştırma Alanları']}")
        doc.add_paragraph() 

    doc.save(word_path)

def get_ders_katalog(url):
    driver.get(url)
    soup = BeautifulSoup(driver.page_source, "html.parser")
    tables = soup.find_all("table", {"width": "100%", "bgcolor": "#FFFFFF"})
    target_table = tables[1]
    rows = target_table.find_all("tr")
    current = ""

    word_path = 'scraping.docx'
    doc = Document(word_path)
    doc.add_heading('Ders Kataloğu', level=1)

    for row in rows:
        # Eğer <tr> satırı dönemi gösteriyorsa, current_period'u güncelle
        if row.find("strong"):
            current = row.get_text(strip=True)
            doc.add_heading(current, level=2)
        else:
            # Dönemi gösteren <tr> değilse, ders bilgilerini çek
            cols = row.find_all("td")
            if len(cols) > 3:
                ders_data = cols[2].get_text(strip=True) + " "+ cols[3].get_text(strip=True) 
                doc.add_paragraph(ders_data)
                doc.add_paragraph("Daha detaylı bilgi için -> ")
                doc.add_paragraph("---------------------------")

    doc.save(word_path)

#get_ogretim_uyeleri("https://ceng.deu.edu.tr/ogretim-uyeleri/")
get_ders_katalog("https://debis.deu.edu.tr/ders-katalog/2024-2025/tr/bolum_1210_tr.html")




